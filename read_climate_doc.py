#!/usr/bin/env python3
"""
Script to read the Big Data-EarthScape Climate Agency document and extract rsr data.
"""

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

def read_climate_document():
    """Read the climate document and extract rsr data."""
    doc_path = Path("Big Data-EarthScape_Climate_Agency.docx")

    if not doc_path.exists():
        print(f"Document not found: {doc_path}")
        return

    try:
        doc = Document(doc_path)
        print("Reading climate document...")
        print("=" * 50)

        rsr_found = False
        rsr_data = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Look for rsr mentions (case insensitive)
                if 'rsr' in text.lower():
                    rsr_found = True
                    rsr_data.append(text)
                    print(f"RSR Data: {text}")
                else:
                    print(f"Content: {text}")

        if not rsr_found:
            print("No RSR data found in the document.")
            print("Searching for related terms...")

            # Look for related climate data terms
            climate_terms = ['climate', 'data', 'agency', 'earthscape', 'big data']
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if any(term in text.lower() for term in climate_terms):
                    print(f"Related: {text}")

        return rsr_data

    except Exception as e:
        print(f"Error reading document: {e}")
        return None

if __name__ == "__main__":
    rsr_data = read_climate_document()
    if rsr_data:
        print(f"\nFound {len(rsr_data)} RSR-related entries")
    else:
        print("\nNo RSR data extracted")
